from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(40, 40, 40)

def add_link(paragraph, url, text):
    from docx.oxml.shared import qn
    from docx.oxml import OxmlElement
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '1A365D')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def create_resume_docx(script_dir):
    doc = Document()
    style_doc(doc)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- HEADER ---
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    name = header.add_run("TONĆI KUČIĆ")
    name.bold = True
    name.font.size = Pt(26)
    name.font.color.rgb = RGBColor(26, 54, 93)
    header.add_run("\n")
    
    tagline = header.add_run("Executive Technology Leader | Board Member | Co-author")
    tagline.bold = True
    tagline.font.size = Pt(13)
    tagline.font.color.rgb = RGBColor(30, 30, 30)
    header.add_run("\n")
    
    contact = header.add_run("Split, Croatia | 003598323033 | tonci.kucic@gmail.com")
    contact.font.size = Pt(10)
    header.add_run("\n")
    
    links_p = doc.add_paragraph()
    links_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_link(links_p, "https://linkedin.com/in/toncikucic", "linkedin.com/in/toncikucic")
    links_p.add_run("  •  ")
    add_link(links_p, "https://tonci77.github.io/tonci-portfolio/", "tonci77.github.io/tonci-portfolio")

    # --- VALUE PROPOSITION ---
    h1 = doc.add_heading('EXECUTIVE PROFILE', level=1)
    h1.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    doc.add_paragraph(
        "Strategic Engineering Executive with over 20 years of experience delivering high-stakes digital transformation "
        "and scaling high-performance organizations. Differentiated by a unique synthesis of deep architectural "
        "mastery, organizational psychology, and boardroom governance. Expert in managing large-scale global "
        "engineering teams (40+) and high-availability BSS/OSS systems in ultra-regulated industries including "
        "iGaming, Telecom, and Life Sciences."
    )

    # --- CORE STRENGTHS & DIFFERENTIATORS ---
    h2 = doc.add_heading('STRATEGIC DIFFERENTIATORS', level=1)
    h2.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    
    diffs = [
        "Human Capital Architecture: Applying organizational psychology frameworks (Shark/Whale/Turtle personas) to scale high-performance cultures and retain top 1% talent.",
        "Precision in Regulation: Veteran lead for Tier-1 operators in mission-critical environments (Telecom/iGaming) where downtime and compliance failure are not options.",
        "Business-First Tech Strategy: Proven ability to translate complex R&D roadmaps into board-level value, P&L optimization, and M&A technical due diligence.",
        "End-to-End System Legacy: From greenfield Solution Architecture for Tier-1 Telcos to leading global engineering for multi-vertical iGaming platforms."
    ]
    
    for d in diffs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(d)

    # --- PROFESSIONAL EXPERIENCE ---
    h3 = doc.add_heading('PROFESSIONAL EXPERIENCE', level=1)
    h3.runs[0].font.color.rgb = RGBColor(26, 54, 93)

    jobs = [
        {
            "title": "Head of Engineering",
            "company": "Casumo (Tier-1 iGaming)",
            "dates": "2024 – Present",
            "desc": "Orchestrating engineering strategy for a global multi-vertical platform. Leading 40+ engineers across Casino, Payments, and Sports. Driving platform resilience and engineering culture using advanced organizational psychology frameworks to ensure high-engagement and technical excellence."
        },
        {
            "title": "Engineering Manager / Team Lead",
            "company": "Studion",
            "dates": "2022 – 2024",
            "desc": "Scalability Architect for a 30+ member engineering organization. Established the Engineering Academy—a benchmark for talent development and internal branding—and implemented rigid engineering standards for distributed teams."
        },
        {
            "title": "Supervisory Board Member",
            "company": "Parkovi i Nasadi d.o.o.",
            "dates": "2023 – 2025",
            "desc": "Providing strategic governance and advisory on digital modernization initiatives and organizational transformation for green infrastructure management systems."
        },
        {
            "title": "Principal Consultant",
            "company": "TONCIK Consulting & Solutions",
            "dates": "2020 – 2025",
            "desc": "Delivering strategic business and technology guidance for Smart City Digitalization, and Intelligent Transportation Systems (ITS) to public and private sector clients."
        },
        {
            "title": "Solution Architect",
            "company": "Hrvatski Telekom",
            "dates": "2020 – 2022",
            "desc": "Lead Architect for greenfield R&D Innovation. Engineered enterprise-scale Product Catalog and Order Capture platforms for Croatia's Tier-1 Telecom provider."
        },
        {
            "title": "Engineering Team Lead",
            "company": "Fortuna Entertainment Group",
            "dates": "2018 – 2020",
            "desc": "Managed engineering delivery for high-load virtual sports platforms across Central Europe. Focused on horizontal scalability, high-availability architecture, and rapid feature velocity."
        },
        {
            "title": "Senior Solution Architect / Lead Developer",
            "company": "Kron d.o.o. (Telecom BSS/OSS)",
            "dates": "2004 – 2018",
            "desc": "14-year progressive leadership role. Designed and integrated mission-critical BSS/OSS systems for Tier-1 Telecom operators. Orchestrated complex system integrations for major telco providers across Southeast Europe."
        }
    ]

    for job in jobs:
        p = doc.add_paragraph()
        p.add_run(f"{job['title']} ").bold = True
        p.add_run(f"| {job['company']}").font.color.rgb = RGBColor(100, 100, 100)
        run = p.add_run(f"\t\t{job['dates']}")
        run.italic = True
        doc.add_paragraph(job['desc']).paragraph_format.space_after = Pt(8)

    # --- TECHNICAL ---
    h4 = doc.add_heading('TECHNICAL & ARCHITECTURAL ECOSYSTEM', level=1)
    h4.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    techs = [
        "Enterprise Architecture: Microservices, Event-Driven, DDD, SOA, System Integration, High-Availability Design.",
        "Cloud & Platform: AWS, GCP, Kubernetes, Docker, Infrastructure as Code (IaC), CI/CD, ELK, Grafana.",
        "Data & Messaging: Apache Kafka, RabbitMQ, PostgreSQL, Oracle, MongoDB, Redis, High-Load Data Management.",
        "Ecosystems: Java/Spring Boot, .NET Core, Python, Node.js, React, Angular."
    ]
    for t in techs:
        doc.add_paragraph(t, style='List Bullet')

    # --- THOUGHT LEADERSHIP ---
    h5 = doc.add_heading('THOUGHT LEADERSHIP & PUBLICATIONS', level=1)
    h5.runs[0].font.color.rgb = RGBColor(26, 54, 93)
    p_pub = doc.add_paragraph()
    p_pub.add_run("Co-author of 'Peak Performance: Mindset & Tools for Managers'").bold = True
    p_pub.add_run("\nA world-class strategic guide to scaling high-performance engineering culture via organizational psychology. ")
    add_link(p_pub, "https://www.amazon.com/Peak-Performance-Mindset-Tools-Managers-ebook/dp/B0DKTV2WBT", "[Purchase Link]")

    docx_path = os.path.join(script_dir, "Tonci_Kucic_Executive_Resume.docx")
    doc.save(docx_path)
    print(f"Resume generated at: {docx_path}")
    return docx_path


def generate_linkedin_guide(script_dir):
    guide_path = os.path.join(script_dir, "LINKEDIN_GUIDE.md")
    content = """# 🚀 LinkedIn Elite Profile Upgrade Guide: Tonći Kučić

This comprehensive guide is engineered to transform your LinkedIn profile into a high-authority **Executive Brand Asset**. It is synchronized with your latest resume data to ensure 100% narrative consistency.

---

## 🎯 OBJECTIVE: The "Boardroom-Ready" Signal
The goal is not just to look "hired," but to look like a **Thought Leader** and **Industry Authority**. Every section must answer three questions for a recruiter or board member:
1.  **Can he scale our organization?** (Scale)
2.  **Does he understand the business?** (Value)
3.  **Is he a modern leader?** (Culture)

---

## 🛠 STEP 1: VISUAL IDENTITY (The First 3 Seconds)

### 1. Banner Image (The "Billboard")
*   **Action:** Do not use the default grey background.
*   **Recommendation:** Use a clean, abstract image of **digital topography, abstract networks, or minimalist architecture** (black/amber theme matches your portfolio).
*   **Alternative:** A photo of you speaking at an event or leading a workshop (high social proof).

### 2. Profile Photo
*   **Status:** Use your current professional headshot.
*   **Tweak:** Ensure it zooms in on your face (smiling/approachable but authoritative). No busy backgrounds.

---

## ✍️ STEP 2: THE HOOK (Headline & Intro)

### 1. The "Power Headline"
This is the most critical text on your profile. It dictates SEO and first impressions.
*   **Copy & Paste This Exact Block:**
    > **Head of Engineering at Casumo | Strategic Technology Executive | Board Member | Co-author of 'Peak Performance' | Scaling High-Performance Cultures through Organizational Psychology**

### 2. Audio Pronunciation
*   **Action:** Use the LinkedIn mobile app to record your name pronunciation.
*   **Bonus:** Add a subtle pitch: *"Tonći Kučić. Scaling engineering excellence."*

---

## 📖 STEP 3: THE NARRATIVE (About Section)

This is your executive pitch. It bridges the gap between "Tech Guy" and "Business Leader."

*   **Copy & Paste This Draft:**

    > **Engineering Executive with 20+ years of experience delivering mission-critical technology solutions and leading global organizational transformations.**
    >
    > I specialize in bridging the gap between complex engineering roadmaps and boardroom-level business strategy. My career has been defined by two pillars: **Technical Precision** in ultra-regulated industries (iGaming, Telecom, Smart Cities) and **Human Capital Architecture**.
    >
    > **🏛 STRATEGIC DIFFERENTIATORS:**
    >
    > *   **The Psychologist of Scale:** As Co-author of *'Peak Performance'*, I apply organizational psychology frameworks (Shark/Whale/Turtle personas) to unlock elite performance in engineering cultures.
    > *   **Precision in Regulation:** Veteran lead for Tier-1 operators (Casumo, Hrvatski Telekom) where downtime and compliance failure are not options.
    > *   **Capital Efficiency:** Proven track record in P&L management, M&A due diligence, and translating R&D into Product-Led Growth.
    >
    > **🚀 CURRENT FOCUS:**
    > Currently leading engineering at **Casumo**, scaling a 40+ person distributed organization to drive platform resilience and multi-vertical growth (Casino, Sports, Payments).
    >
    > **📚 AUTHORSHIP:**
    > Co-author of **"Peak Performance: Mindset & Tools for Managers"** — A guide for modern leaders on building resilient, high-efficacy teams.
    >
    > **🌐 DIGITAL ECOSYSTEM:**
    > Detailed Case Studies & Portfolio: https://tonci77.github.io/tonci-portfolio/

---

## 💼 STEP 4: EXPERIENCE (The Evidence)

Don't just list tasks. List **Impact** and **Scale**.

### 1. Casumo (Head of Engineering)
*   **Action:** Update your current role.
*   **Description:**
    > **Strategic Ownership:** Directing engineering strategy for a Tier-1 global iGaming platform.
    > **Scale:** Leading 40+ engineers across Payments, Casino, and Sports verticals.
    > **Impact:**
    > *   Driving platform resilience and engineering culture transformation.
    > *   Integrated internal frameworks for engineering excellence and talent density.
    > *   Architecting high-availability systems for global scale in regulated markets.

### 2. Parkovi i Nasadi (Supervisory Board Member)
*   **Action:** Ensure this is listed to signal **Governance Experience**.
*   **Description:**
    > Providing strategic governance and advisory on digital modernization initiatives and organizational transformation for green infrastructure management systems.

### 3. TONCIK Consulting (Principal Consultant)
*   **Action:** Highlight your advisory capacity.
*   **Description:**
    > Delivering strategic business and technology guidance for Smart City Digitalization and Intelligent Transportation Systems (ITS) to public and private sector clients.

---

## 🌟 STEP 5: FEATURED SECTION (The Halo Effect)

This is the big visual carousel at the top of your profile. **Strategic placement is key.**

1.  **Slot 1 (The Portfolio):** Link to `https://tonci77.github.io/tonci-portfolio/`. Use a thumbnail of your website's Hero section.
2.  **Slot 2 (The Book):** Link to your Amazon book page. Title it: *"Co-authored Book: Peak Performance"*.
3.  **Slot 3 (The Resume):** Upload the PDF generated in this folder (`Tonci_Kucic_Executive_Resume.pdf`). Title it: *"Executive Resume 2026"*.

---

## 💡 STEP 6: SKILLS & ENDORSEMENTS (SEO Optimization)

LinkedIn's algorithm searches for keywords here. Ensure you have 50 skills listed.
*   **Top 3 Constraints (Pin These):**
    1.  **Technology Strategy**
    2.  **Organizational Psychology**
    3.  **Enterprise Architecture**
*   **Other Critical Keywords:**
    *   Executive Leadership
    *   Digital Transformation
    *   iGaming
    *   Telecommunications
    *   Distributed Team Management
    *   Product-Led Growth
    *   Board Advisory

---

*Last Updated for Executive Impact: 2026-01-28*
"""
    with open(guide_path, "w") as f:
        f.write(content)
    print(f"LinkedIn Guide generated at: {guide_path}")


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    # 1. Generate DOCX
    docx_path = create_resume_docx(script_dir)
    
    # 2. Generate PDF
    pdf_path = os.path.join(script_dir, "Tonci_Kucic_Executive_Resume.pdf")
    try:
        from docx2pdf import convert
        print("Converting to PDF (requires Microsoft Word to be open)...")
        convert(docx_path, pdf_path)
        print(f"PDF generated at: {pdf_path}")
    except Exception as e:
        print(f"\n[Note] PDF conversion skipped or failed: {e}")
        print("PDF conversion on Mac requires Microsoft Word to be installed.")
    
    # 3. Generate LinkedIn Guide
    generate_linkedin_guide(script_dir)
